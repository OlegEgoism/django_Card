import datetime
from io import BytesIO

from django.contrib.auth import logout, login
from django.http import HttpResponse
from django.shortcuts import render, redirect, get_object_or_404
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
from docx.shared import Mm, Cm, Pt
from card_personnel.forms import PersonelForm, ContractForm, RegistrationForm, EducationForm, LanguageSkillsForm, WorkForm, FamilyForm, PhoneForm, LoginForm, RegistrationUserForm
from card_personnel.models import Personel, Contract, Registration, Education, LanguageSkills, Family, Phone, Work


def user_register(request):
    if request.method == 'POST':
        form = RegistrationUserForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect('personel_list')
    else:
        form = RegistrationUserForm()
    context = {
        'form': form
    }
    return render(request, 'register.html', context)


def user_login(request):
    if request.method == 'POST':
        form = LoginForm(data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            return redirect('personel_list')
    else:
        form = LoginForm()
    context = {
        'form': form
    }
    return render(request, 'login.html', context)


def user_logout(request):
    logout(request)
    return redirect('login')


def personel_list(request):
    personel_list = Personel.objects.all()
    context = {
        'personel_list': personel_list
    }
    return render(request, 'personel_list.html', context)


def personel_detail(request, id):
    personel = Personel.objects.get(id=id)
    context = {
        'personel': personel
    }
    return render(request, 'personel_detail.html', context)


def personel_create(request):
    if request.method == 'POST':
        personel_form = PersonelForm(request.POST, request.FILES)
        contract_form = ContractForm(request.POST)
        registration_form = RegistrationForm(request.POST)
        if personel_form.is_valid() and contract_form.is_valid() and registration_form.is_valid():
            personel = personel_form.save()
            contract = contract_form.save(commit=False)
            contract.personel = personel
            contract.save()
            registration_form.instance.personel = personel
            registration_form.save()
            return redirect('personel_list')
    else:
        personel_form = PersonelForm()
        contract_form = ContractForm()
        registration_form = RegistrationForm()
    context = {
        'personel_form': personel_form,
        'registration_form': registration_form,
        'contract_form': contract_form,
    }
    return render(request, 'personel_create.html', context)


def personel_delete(request, id):
    try:
        personel = Personel.objects.get(pk=id)
        personel.delete()
        return redirect('personel_list')
    except Personel.DoesNotExist:
        return redirect('personel_list')


def personel_edit(request, id):
    personel = get_object_or_404(Personel, id=id)
    contract = Contract.objects.get(personel=personel)
    registration = Registration.objects.get(personel=personel)
    if request.method == 'POST':
        personel_form = PersonelForm(request.POST, request.FILES, instance=personel)
        contract_form = ContractForm(request.POST, instance=contract)
        registration_form = RegistrationForm(request.POST, instance=registration)
        if personel_form.is_valid() and contract_form.is_valid() and registration_form.is_valid():
            personel = personel_form.save()
            if not contract:
                contract = contract_form.save(commit=False)
                contract.personel = personel
                contract.save()
            else:
                contract_form.save()
            if not registration:
                registration = registration_form.save(commit=False)
                registration.personel = personel
                registration.save()
            else:
                registration_form.save()
            return redirect('personel_detail', id=id)
    else:
        personel_form = PersonelForm(instance=personel)
        contract_form = ContractForm(instance=contract) if contract else ContractForm()
        registration_form = RegistrationForm(instance=registration) if registration else RegistrationForm()
    context = {
        'personel_form': personel_form,
        'registration_form': registration_form,
        'contract_form': contract_form,
        'personel': personel,
    }
    return render(request, 'personel_edit.html', context)


def phone_create(request, id):
    personel = Personel.objects.get(id=id)
    if request.method == 'POST':
        phone_form = PhoneForm(request.POST)
        if phone_form.is_valid():
            phone = phone_form.save(commit=False)
            phone.personel = personel
            phone.save()
            return redirect('personel_detail', id=id)
    else:
        phone_form = PhoneForm()
    context = {
        'personel': personel,
        'phone_form': phone_form,
    }
    return render(request, 'phone_create.html', context)


def phone_delete(request, id, phone_id):
    personel = get_object_or_404(Personel, id=id)
    phone = get_object_or_404(Phone, id=phone_id)
    if request.method == 'POST':
        phone.delete()
        return redirect('personel_detail', id=id)
    context = {
        'personel': personel,
        'phone': phone,
    }
    return render(request, 'delete.html', context)


def phone_edit(request, phone_id):
    phone = get_object_or_404(Phone, id=phone_id)
    if request.method == 'POST':
        form = PhoneForm(request.POST, instance=phone)
        if form.is_valid():
            form.save()
            return redirect('personel_detail', id=phone.personel.id)
    else:
        form = PhoneForm(instance=phone)
    context = {
        'form': form,
        'phone': phone,
    }
    return render(request, 'edit.html', context)


def work_create(request, id):
    personel = Personel.objects.get(id=id)
    if request.method == 'POST':
        work_form = WorkForm(request.POST)
        if work_form.is_valid():
            work = work_form.save(commit=False)
            work.personel = personel
            work.save()
            return redirect('personel_detail', id=id)
    else:
        work_form = WorkForm()
    context = {
        'personel': personel,
        'work_form': work_form,
    }
    return render(request, 'work_create.html', context)


def work_delete(request, id, work_id):
    personel = get_object_or_404(Personel, id=id)
    work = get_object_or_404(Work, id=work_id)
    if request.method == 'POST':
        work.delete()
        return redirect('personel_detail', id=id)
    context = {
        'personel': personel,
        'work': work,
    }
    return render(request, 'delete.html', context)


def work_edit(request, work_id):
    work = get_object_or_404(Work, id=work_id)
    if request.method == 'POST':
        form = WorkForm(request.POST, instance=work)
        if form.is_valid():
            form.save()
            return redirect('personel_detail', id=work.personel.id)
    else:
        form = WorkForm(instance=work)
    context = {
        'form': form,
        'work': work,
    }
    return render(request, 'edit.html', context)


def education_create(request, id):
    personel = Personel.objects.get(id=id)
    if request.method == 'POST':
        education_form = EducationForm(request.POST)
        if education_form.is_valid():
            education = education_form.save(commit=False)
            education.personel = personel
            education.save()
            return redirect('personel_detail', id=id)
    else:
        education_form = EducationForm()
    context = {
        'personel': personel,
        'education_form': education_form,
    }
    return render(request, 'education_create.html', context)


def education_delete(request, id, education_id):
    personel = get_object_or_404(Personel, id=id)
    education = get_object_or_404(Education, id=education_id)
    if request.method == 'POST':
        education.delete()
        return redirect('personel_detail', id=id)
    context = {
        'personel': personel,
        'education': education,
    }
    return render(request, 'delete.html', context)


def education_edit(request, education_id):
    education = get_object_or_404(Education, id=education_id)
    if request.method == 'POST':
        form = EducationForm(request.POST, instance=education)
        if form.is_valid():
            form.save()
            return redirect('personel_detail', id=education.personel.id)
    else:
        form = EducationForm(instance=education)
    context = {
        'form': form,
        'education': education,
    }
    return render(request, 'edit.html', context)


def language_skill_create(request, id):
    personel = Personel.objects.get(id=id)
    if request.method == 'POST':
        language_skill_form = LanguageSkillsForm(request.POST)
        if language_skill_form.is_valid():
            language_skill = language_skill_form.save(commit=False)
            language_skill.personel = personel
            language_skill.save()
            return redirect('personel_detail', id=id)
    else:
        language_skill_form = LanguageSkillsForm()
    context = {
        'personel': personel,
        'language_skill_form': language_skill_form,
    }
    return render(request, 'language_skill_create.html', context)


def language_skill_delete(request, id, language_skill_id):
    personel = get_object_or_404(Personel, id=id)
    language_skill = get_object_or_404(LanguageSkills, id=language_skill_id)
    if request.method == 'POST':
        language_skill.delete()
        return redirect('personel_detail', id=id)
    context = {
        'personel': personel,
        'language_skill': language_skill,
    }
    return render(request, 'delete.html', context)


def language_skill_edit(request, language_skill_id):
    language_skill = get_object_or_404(LanguageSkills, id=language_skill_id)
    if request.method == 'POST':
        form = LanguageSkillsForm(request.POST, instance=language_skill)
        if form.is_valid():
            form.save()
            return redirect('personel_detail', id=language_skill.personel.id)
    else:
        form = LanguageSkillsForm(instance=language_skill)
    context = {
        'form': form,
        'language_skill': language_skill,
    }
    return render(request, 'edit.html', context)


def family_create(request, id):
    personel = Personel.objects.get(id=id)
    if request.method == 'POST':
        family_form = FamilyForm(request.POST)
        if family_form.is_valid():
            family = family_form.save(commit=False)
            family.personel = personel
            family.save()
            return redirect('personel_detail', id=id)
    else:
        family_form = FamilyForm()
    context = {
        'personel': personel,
        'family_form': family_form,
    }
    return render(request, 'family_create.html', context)


def family_delete(request, id, family_id):
    personel = get_object_or_404(Personel, id=id)
    family = get_object_or_404(Family, id=family_id)
    if request.method == 'POST':
        family.delete()
        return redirect('personel_detail', id=id)
    context = {
        'personel': personel,
        'family': family,
    }
    return render(request, 'delete.html', context)


def family_edit(request, family_id):
    family = get_object_or_404(Family, id=family_id)
    if request.method == 'POST':
        form = FamilyForm(request.POST, instance=family)
        if form.is_valid():
            form.save()
            return redirect('personel_detail', id=family.personel.id)
    else:
        form = FamilyForm(instance=family)
    context = {
        'form': form,
        'family': family,
    }
    return render(request, 'edit.html', context)


def spravka(request, id):
    personel = None
    contract = None
    registration = None
    try:
        personel = Personel.objects.get(pk=id)
        contract = Contract.objects.get(personel=personel)
        registration = Registration.objects.get(personel=personel)
    except Personel.DoesNotExist:
        personel = None
    except Contract.DoesNotExist:
        contract = None
    except Registration.DoesNotExist:
        registration = None

    """Создаем документ"""
    doc = Document()
    section = doc.sections[0]
    # Формат страницы
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    # Установите размеры страницы
    section.left_margin = Mm(20)
    section.right_margin = Mm(15)
    section.top_margin = Mm(12)
    section.bottom_margin = Mm(8)

    """СПРАВКА"""
    paragraph = doc.add_paragraph('СПРАВКА')
    run = paragraph.runs[0]
    run.font.size = Pt(14)  # Установка размера шрифта
    run.font.name = 'Times New Roman'  # Установка стиля шрифта
    run.bold = True  # Установка полужирного стиля шрифта
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Выравнивание по центру

    """ФИО"""
    paragraph = doc.add_paragraph(f'{personel.last_name.upper()} {personel.first_name.upper()} {personel.middle_name.upper()}')
    run = paragraph.runs[0]
    run.font.size = Pt(12)  # Установка размера шрифта
    run.font.name = 'Times New Roman'  # Установка стиля шрифта
    run.bold = True  # Установка полужирного стиля шрифта
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Выравнивание по центру

    """ТАБЛИЦА 1"""
    table = doc.add_table(rows=5, cols=2)
    table.autofit = False
    for cell in table.columns[0].cells:
        cell.width = Cm(12)
    for cell in table.columns[1].cells:
        cell.width = Cm(6)

    table.cell(0, 1).merge(table.cell(4, 1))

    cell = table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run('Фамилия: ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run = paragraph.add_run(personel.last_name)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.underline = True
    run.underline = WD_UNDERLINE.SINGLE

    cell = table.cell(1, 0)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run('Имя: ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run = paragraph.add_run(personel.first_name)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.underline = True
    run.underline = WD_UNDERLINE.SINGLE

    cell = table.cell(2, 0)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run('Отчество: ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run = paragraph.add_run(personel.middle_name)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.underline = True
    run.underline = WD_UNDERLINE.SINGLE

    cell = table.cell(3, 0)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run('Пол: ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run = paragraph.add_run(personel.gender.name)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.underline = True
    run.underline = WD_UNDERLINE.SINGLE

    cell = table.cell(4, 0)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run('Дата рождения: ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run = paragraph.add_run(f'{personel.date_of_birth.strftime("%d.%m.%Y")} г.')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.underline = True
    run.underline = WD_UNDERLINE.SINGLE

    if personel.photo and personel.photo.file:
        image_stream = BytesIO(personel.photo.read())
        cell = table.cell(0, 1)
        cell.paragraphs[0].add_run().add_picture(image_stream, width=Mm(30), height=Mm(40))
    else:
        cell = table.cell(0, 1)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run('Место для фотографии в сканированном цветном изображении либо наклеенном виде (30х40 мм)')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

    if contract:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run('Контракт: ')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        run = paragraph.add_run(f'{contract.type} ')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.underline = True
        run.underline = WD_UNDERLINE.SINGLE

        run = paragraph.add_run(f'c {contract.date_start.strftime("%d.%m.%Y")} г. ')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.underline = True
        run.underline = WD_UNDERLINE.SINGLE

        run = paragraph.add_run(f'по {contract.date_end.strftime("%d.%m.%Y")} г.')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.underline = True
        run.underline = WD_UNDERLINE.SINGLE
    else:
        paragraph = doc.add_paragraph('Контракт: Нет информации')
        run = paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.underline = True
        run.underline = WD_UNDERLINE.SINGLE

    if registration:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run('Место жительства: ')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        run = paragraph.add_run(f'{registration.country.name}, ')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.underline = True
        run.underline = WD_UNDERLINE.SINGLE

        run = paragraph.add_run(f'р-н. {registration.district}, ')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.underline = True
        run.underline = WD_UNDERLINE.SINGLE

        run = paragraph.add_run(f'г. {registration.city}, ')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.underline = True
        run.underline = WD_UNDERLINE.SINGLE

        run = paragraph.add_run(f'ул. {registration.street} ')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.underline = True
        run.underline = WD_UNDERLINE.SINGLE

        run = paragraph.add_run(f'г. {registration.house}, ')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.underline = True
        run.underline = WD_UNDERLINE.SINGLE

        run = paragraph.add_run(f'кв. {registration.apartment}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.underline = True
        run.underline = WD_UNDERLINE.SINGLE
    else:
        paragraph = doc.add_paragraph('Место жительства: Нет информации')
        run = paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.underline = True
        run.underline = WD_UNDERLINE.SINGLE

    phone = Phone.objects.filter(personel=personel)
    phone_info_str = ''
    if phone:
        phone_info_list = [f"{pn.number}" for pn in phone]
        phone_info_str = ', '.join(phone_info_list)
    else:
        phone_info_str = 'Нет информации'
    new_paragraph = doc.add_paragraph()
    run = new_paragraph.add_run('Номер телефона: ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run = new_paragraph.add_run(phone_info_str)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.underline = True
    run.underline = WD_UNDERLINE.SINGLE

    new_paragraph = doc.add_paragraph()
    new_paragraph.paragraph_format.space_after = Pt(4)
    run = new_paragraph.add_run('Образование.')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    """ТАБЛИЦА 2"""
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    for cell in table.columns[0].cells:
        cell.width = Cm(7)
    for cell in table.columns[1].cells:
        cell.width = Cm(2)
    for cell in table.columns[2].cells:
        cell.width = Cm(1)
    for cell in table.columns[3].cells:
        cell.width = Cm(1)
    for cell in table.columns[4].cells:
        cell.width = Cm(5)
    column_headers = ['Название, адрес учебного заведения',
                      'Уровень образования',
                      'Дата начала',
                      'Дата окончания',
                      'Специальность'
                      ]
    for i, header in enumerate(column_headers):
        cell = table.cell(0, i)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        paragraph.spacing_after = Pt(0)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    education = Education.objects.filter(personel=personel)
    for education in education:
        row = table.add_row().cells
        if education.name:
            education_info = education.name
            if education.address:
                education_info += '. ' + education.address
            row[0].paragraphs[0].add_run(education_info).font.name = 'Times New Roman'
            row[0].paragraphs[0].runs[0].font.size = Pt(12)
        else:
            row[0].paragraphs[0].add_run('').font.name = 'Times New Roman'
            row[0].paragraphs[0].runs[0].font.size = Pt(12)
        if education.level:
            row[1].paragraphs[0].add_run(education.level.level).font.name = 'Times New Roman'
            row[1].paragraphs[0].runs[0].font.size = Pt(12)
        else:
            row[1].paragraphs[0].add_run('').font.name = 'Times New Roman'
            row[1].paragraphs[0].runs[0].font.size = Pt(12)
        if education.date_start:
            row[2].paragraphs[0].add_run(str(education.date_start)).font.name = 'Times New Roman'
            row[2].paragraphs[0].runs[0].font.size = Pt(12)
        else:
            row[2].paragraphs[0].add_run('').font.name = 'Times New Roman'
            row[2].paragraphs[0].runs[0].font.size = Pt(12)
        if education.date_end:
            row[3].paragraphs[0].add_run(str(education.date_end)).font.name = 'Times New Roman'
            row[3].paragraphs[0].runs[0].font.size = Pt(12)
        else:
            row[3].paragraphs[0].add_run('').font.name = 'Times New Roman'
            row[3].paragraphs[0].runs[0].font.size = Pt(12)
        if education.speciality:
            row[4].paragraphs[0].add_run(education.speciality).font.name = 'Times New Roman'
            row[4].paragraphs[0].runs[0].font.size = Pt(12)
        else:
            row[4].paragraphs[0].add_run('').font.name = 'Times New Roman'
            row[4].paragraphs[0].runs[0].font.size = Pt(12)
    if not education:
        row = table.add_row().cells
        for i in range(5):
            cell = row[i]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run('-')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    language_skills = LanguageSkills.objects.filter(personel=personel)
    language_skills_str = ''
    if language_skills:
        language_skills_list = [f"{ls.language.name} ({ls.level.level})" for ls in language_skills]
        language_skills_str = ', '.join(language_skills_list)
    else:
        language_skills_str = 'Нет информации'
    new_paragraph = doc.add_paragraph()
    new_paragraph.paragraph_format.space_before = Pt(12)
    run = new_paragraph.add_run('Владение языками: ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run = new_paragraph.add_run(language_skills_str)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.underline = True
    run.underline = WD_UNDERLINE.SINGLE

    family = Family.objects.filter(personel=personel)
    family_info_list = []
    for member in family:
        family_info = f'{member.last_name} {member.first_name} {member.middle_name}, {member.status.name}, {member.date_of_birth.strftime("%d.%m.%Y")} г.р.'
        family_info_list.append(family_info)
    if family_info_list:
        family_info_str = ', '.join(family_info_list)
    else:
        family_info_str = 'Нет информации'
    new_paragraph = doc.add_paragraph()
    run = new_paragraph.add_run('Состав семьи: ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run = new_paragraph.add_run(family_info_str)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.underline = True
    run.underline = WD_UNDERLINE.SINGLE

    """ДВА ПРОБЕЛА"""
    paragraph = doc.add_paragraph(' ')
    run = paragraph.runs[0]
    run.font.size = Pt(14)  # Установка размера шрифта
    run.font.name = 'Times New Roman'  # Установка стиля шрифта
    paragraph = doc.add_paragraph(' ')
    run = paragraph.runs[0]
    run.font.size = Pt(14)  # Установка размера шрифта
    run.font.name = 'Times New Roman'  # Установка стиля шрифта

    """ДАТА ЗАПОЛНЕНИЯ"""
    current_date = datetime.datetime.now().strftime("%d.%m.%Y")
    paragraph = doc.add_paragraph(current_date)
    run = paragraph.runs[0]
    run.font.size = Pt(12)  # Установка размера шрифта
    run.font.name = 'Times New Roman'  # Установка стиля шрифта

    """ПОДГОТОВКА ДЛЯ СОХРАНЕНИЯ"""
    output = BytesIO()
    doc.save(output)
    response = HttpResponse(content_type='application/msword')
    response['Content-Disposition'] = 'attachment; filename="справка.docx"'
    response.write(output.getvalue())
    return response
